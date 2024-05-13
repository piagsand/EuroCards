# /api/generate_doc.py
import json
from http.server import BaseHTTPRequestHandler
from docx import Document
import io

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        data = json.loads(post_data.decode('utf-8'))

        doc = Document()
        doc.add_heading('Liste des Cartes', level=1)

        sections = [
            ('Cartes manquantes', data['missingCards']),
            ('Cartes acquises', data['ownedCards']),
            ('Cartes doubles', data['duplicateCards'])
        ]

        for title, cards in sections:
            doc.add_heading(title, level=2)
            for card in cards:
                doc.add_paragraph(f"{card}")

        # Sauvegarde du document en mémoire
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)

        self.send_response(200)
        self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.end_headers()
        self.wfile.write(mem_file.getvalue())

