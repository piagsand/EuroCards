# -*- coding: utf-8 -*-

import json
import sys
print(sys.path)
from collections import defaultdict
from docx import Document




def format_list_by_country(card_list, country_codes):
    """Groupe et formate les cartes par préfixe de pays et ajoute les cartes non reconnues à une catégorie spéciale."""
    country_dict = defaultdict(list)
    special_cards = []

    for card in card_list:
        prefix = card.split('-')[0]
        if prefix in country_codes:
            country_dict[country_codes[prefix]].append(card)
        else:
            special_cards.append(card)

    if special_cards:
        country_dict["Cartes spéciales"] = special_cards

    formatted_list = []
    for country, cards in sorted(country_dict.items(), key=lambda x: x[0] != "Cartes spéciales"):
        formatted_cards = " / ".join(cards)
        formatted_list.append((country, formatted_cards))  # Assurez-vous que ceci est un tuple de deux éléments

    return formatted_list

# Chemin vers le fichier JSON original
file_path = 'cards.json'

# Charger les données JSON à partir du fichier original
with open(file_path, 'r') as file:
    data = json.load(file)

country_codes = {
    "GER": "Germany",
    "SCO": "Scotland",
    "HUN": "Hungary",
    "SUI": "Switzerland",
    "ES": "Spain",
    "CRO": "Croatia",
    "ITA": "Italy",
    "ALB": "Albania",
    "SVN": "Slovenia",
    "DEN": "Denmark",
    "SRB": "Serbia",
    "ENG": "England",
    "POL": "Poland",
    "EST": "Estonia",
    "WAL": "Wales",
    "LEG": "Euro Legends",
    "FIN": "Finland",
    "NED": "Netherlands",
    "AUT": "Austria",
    "FRA": "France",
    "BEL": "Belgium",
    "SVK": "Slovakia",
    "ROM": "Romania",
    "ISR": "Israel",
    "ICE": "Iceland",
    "BIH": "Bosnia and Herzegovina",
    "UKR": "Ukraine",
    "TUR": "Turkiye",
    "GEO": "Georgia",
    "LUX": "Luxembourg",
    "GREC": "Greece",
    "KAZ": "Kazakhstan",
    "POR": "Portugal",
    "CZE": "Czechia"
}


formatted_missing = format_list_by_country(data.get('missingCards', []), country_codes)
formatted_owned = format_list_by_country(data.get('ownedCards', []), country_codes)
formatted_doubles = format_list_by_country(data.get('doubleCards', []), country_codes)

# Création d'un document Word
doc = Document()
doc.add_heading('Liste des Cartes', level=1)

# Ajout des sections avec les cartes
sections = [('Cartes manquantes', formatted_missing), ('Cartes acquises', formatted_owned), ('Cartes doubles', formatted_doubles)]
for title, content in sections:
    doc.add_heading(title, level=2)
    for country, cards in content:  # Ici, chaque 'content' doit être une liste de tuples (country, cards)
        doc.add_paragraph("{} : {}".format(country, cards))


# Sauvegarde du document
output_file_path = 'Cartes_Report.docx'
doc.save(output_file_path)

print("Le document Word a été créé avec succès et est enregistré sous le nom '{}'".format(output_file_path))
