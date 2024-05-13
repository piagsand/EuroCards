# -*- coding: utf-8 -*-

from http.server import BaseHTTPRequestHandler, HTTPServer
from docx import Document
import io
import json

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
        except (ValueError, json.JSONDecodeError):
            self.send_response(400)  # Bad Request
            self.end_headers()
            return

        # Country codes mapping
        country_codes = {
            "GER": "Germany", "SCO": "Scotland", "HUN": "Hungary", "SUI": "Switzerland",
            "ESP": "Spain", "CRO": "Croatia", "ITA": "Italy", "ALB": "Albania", "SVN": "Slovenia",
            "DEN": "Denmark", "SRB": "Serbia", "ENG": "England", "POL": "Poland", "EST": "Estonia",
            "WAL": "Wales", "FIN": "Finland", "NED": "Netherlands", "AUT": "Austria", "FRA": "France",
            "BEL": "Belgium", "SVK": "Slovakia", "ROM": "Romania", "ISR": "Israel", "ICE": "Iceland",
            "BIH": "Bosnia and Herzegovina", "UKR": "Ukraine", "TUR": "Turkey", "GEO": "Georgia",
            "LUX": "Luxembourg", "GRE": "Greece", "KAZ": "Kazakhstan", "POR": "Portugal", "CZE": "Czechia"
        }

        # Create a Word document
        doc = Document()
        doc.add_heading('Liste des Cartes', level=1)

        # Process different card categories
        sections = [
            ('Cartes manquantes', data.get('missingCards', [])),
            ('Cartes acquises', data.get('ownedCards', [])),
            ('Cartes doubles', data.get('doubleCards', []))
        ]
        for title, cards in sections:
            doc.add_heading(title, level=2)
            for card in cards:
                prefix = card.split('-')[0]
                country_name = country_codes.get(prefix, 'Unknown Country')
                doc.add_paragraph(f"{country_name}: {card}")

        # Save the Word document in memory
        with io.BytesIO() as mem_file:
            doc.save(mem_file)
            mem_file.seek(0)
            document_data = mem_file.getvalue()
            self.send_response(200)
            self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Length', str(len(document_data)))
            self.end_headers()
            self.wfile.write(document_data)

    def do_OPTIONS(self):
        self.send_response(204)  # No Content
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

def run(server_class=HTTPServer, handler_class=handler, port=8080):
    server_address = ('', port)
    httpd = server_class(server_address, handler_class)
    print(f'Starting httpd on port {port}...')
    httpd.serve_forever()

if __name__ == "__main__":
    run()